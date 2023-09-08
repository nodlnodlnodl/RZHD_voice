import json
from docx import Document

# Загрузите docx документ
doc = Document("documentation_train.docx")


# Функция для извлечения названия электропоезда
def extract_train_name(paragraph):
    for run in paragraph.runs:
        if run.bold and "Порядок действий локомотивных бригад" in run.text:
            for word in run.text.split():
                if any(char.isdigit() for char in word) and any(char.isupper() for char in word):
                    return word
    return None


data = []
train_name = None
id_counter = 1

paragraph_index = 0
table_index = 0
heading1_str = '---'
while paragraph_index < len(doc.paragraphs):
    paragraph = doc.paragraphs[paragraph_index]
    train_name_candidate = extract_train_name(paragraph)

    if train_name_candidate:
        train_name = train_name_candidate
        paragraph_index += 1

        if table_index < len(doc.tables):
            table = doc.tables[table_index]
            for row in table.rows:
                if len(row.cells) > 2 and row.cells[1].text == row.cells[2].text:
                    heading1_str = row.cells[1].text
                heading2_str = row.cells[1].text if len(row.cells) > 1 and row.cells[1].text != 'Неисправность' and row.cells[1].text != row.cells[2].text else "---"
                heading3_str = row.cells[2].text if len(row.cells) > 2 and row.cells[2].text != 'Вероятная причина' and row.cells[1].text != row.cells[2].text else "---"
                text_to_do = row.cells[3].text if len(row.cells) > 3 and row.cells[3].text != ('Метод устранения' or 'Метод устранения неисправности') and row.cells[2].text != row.cells[3].text else "---"
                heading1_str = heading1_str.replace('\n', ' ')
                heading2_str = heading2_str.replace('\n', ' ')
                heading3_str = heading3_str.replace('\n', ' ')
                text_to_do = text_to_do.replace('\n', ' ')
                if text_to_do != '---':
                    entry = {
                        "id": id_counter,
                        "title": train_name,
                        "heading1": heading1_str,
                        "heading2": heading2_str,
                        "heading3": heading3_str,
                        "heading4": "---",
                        "heading5": "---",
                        "text": text_to_do
                    }
                    data.append(entry)
                    id_counter += 1
                    print(entry)
            table_index += 1
    else:
        paragraph_index += 1

# Сохраните данные в формате JSON
with open('kb.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, indent=4)

print("Готово!")
